#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fill_docx.py — docx 合并单元格表格定位与填充 (application-bundle skill)
用法:
  python fill_docx.py dump <模板.docx>             # 打印每个表格的合并单元格结构
  python fill_docx.py fill <模板.docx> <映射.csv> <输出.docx>  # 按 "表,行,列,文本" 填充
只读输入；fill 写新文件，不覆盖模板。
"""
import sys, csv, os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def dump_structure(path):
    d = Document(path)
    for ti, t in enumerate(d.tables):
        print(f"===== 表{ti} ({len(t.rows)}x{len(t.columns)}) =====")
        for ri, row in enumerate(t.rows):
            seen = {}
            parts = []
            for ci, cell in enumerate(row.cells):
                key = id(cell._tc)
                if key in seen:
                    parts.append(f"[{seen[key]}<-合并]")
                else:
                    seen[key] = ci
                    parts.append(f"({ci}:{cell.text.strip()[:10]})")
            print(f"R{ri}: " + " ".join(parts))

def set_cell(cell, text, size=12, keep_tail=None):
    lines = text.split("\n")
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    for i, line in enumerate(lines):
        p = p0 if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
    if keep_tail:
        for tl in keep_tail.split("\n"):
            p = cell.add_paragraph()
            run = p.add_run(tl)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(size)

def fill(path, mapping_csv, out):
    d = Document(path)
    rows = list(csv.reader(open(mapping_csv, encoding="utf-8")))
    for r in rows:
        if len(r) < 4:
            continue
        t_idx, r_idx, c_idx, text = int(r[0]), int(r[1]), int(r[2]), r[3]
        keep = r[4] if len(r) > 4 and r[4] else None
        cell = d.tables[t_idx].rows[r_idx].cells[c_idx]
        set_cell(cell, text, keep_tail=keep)
    d.save(out)
    print("saved:", out)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)
    cmd, path = sys.argv[1], sys.argv[2]
    if cmd == "dump":
        dump_structure(path)
    elif cmd == "fill":
        fill(path, sys.argv[3], sys.argv[4])
    else:
        print("未知命令"); sys.exit(2)
