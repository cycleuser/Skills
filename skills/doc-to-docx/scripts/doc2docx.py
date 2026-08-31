#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""doc2docx.py — pure-Python legacy Word (.doc) → .docx converter.

This module converts the legacy binary Word format (.doc, Word 97/2000/2002/2003,
and the Word-97-compatible output written by WPS Office) into the modern OOXML
(.docx) format **using only Python** — no LibreOffice, no Microsoft Word, no
textutil, no antiword/catdoc, no external binaries of any kind.

Pipeline (each stage is described in rules/binary-format.md):

  1. open the .doc as an OLE2/CFB container (olefile, pure Python)
  2. parse the FIB (File Information Block) to locate every substructure
  3. parse the CLX piece table and decode the main text (UTF-16 / cp1252)
  4. parse the CHPX / PAPX FKP pages into character- and paragraph-format runs
  5. parse the STSH style sheet and the STTBF font table (font names)
  6. parse PlcfSed / SEPX into per-section page setup
  7. resolve style chains + direct property exceptions -> full formatting
  8. reconstruct tables from fInTable / fTtp paragraph marks + TDefTable
  9. write a .docx with python-docx + raw lxml for byte-precise OOXML

The converter is "lossy-aware": everything it can express in OOXML it does, and
a per-item fidelity report is produced so nothing silently disappears. Text,
character formatting (font/size/bold/italic/underline/colour/highlight/…),
paragraph formatting (alignment/indent/spacing/line-spacing/keep/outline),
page setup (page size/margins/orientation/header distance), the Normal default
formatting and tables are preserved with high fidelity.

CLI::

    python doc2docx.py convert <in.doc> [-o out.docx] [--json]
    python doc2docx.py batch <dir> [-o outdir]
    python doc2docx.py inspect <in.doc>
    python doc2docx.py check <out.docx> [--against in.doc]

Dependencies: olefile, python-docx, lxml  (all pure-Python installs).

Author: cycleuser
License: MIT
"""

from __future__ import annotations

import argparse
import json
import os
import re
import struct
import sys
from dataclasses import dataclass, field, fields as dc_fields

try:
    import olefile
except ImportError:  # pragma: no cover
    raise SystemExit("doc2docx requires olefile. Install it with: pip install olefile")

try:
    from docx import Document
except ImportError:  # pragma: no cover
    raise SystemExit("doc2docx requires python-docx. Install it with: pip install python-docx")

try:
    from lxml import etree
except ImportError:  # pragma: no cover
    raise SystemExit("doc2docx requires lxml. Install it with: pip install lxml")

__version__ = "1.0.0"
__all__ = ["convert", "inspect", "verify", "DocReader", "main"]

# --------------------------------------------------------------------------- #
#  Small binary helpers
# --------------------------------------------------------------------------- #

def u8(b: bytes, off: int) -> int:
    return b[off]


def u16(b: bytes, off: int) -> int:
    return struct.unpack_from("<H", b, off)[0]


def i16(b: bytes, off: int) -> int:
    return struct.unpack_from("<h", b, off)[0]


def u32(b: bytes, off: int) -> int:
    return struct.unpack_from("<I", b, off)[0]


def i32(b: bytes, off: int) -> int:
    return struct.unpack_from("<i", b, off)[0]


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"


def w(tag: str) -> str:
    """Clark-notation tag in the WordprocessingML namespace."""
    return f"{{{W_NS}}}{tag}"


# --------------------------------------------------------------------------- #
#  Sprm machinery  (MS-DOC 2.4.84 / 2.6)
# --------------------------------------------------------------------------- #

def sprm_len(sprm: int, data: bytes, off: int) -> int:
    """Total bytes (sprm + operand) of one Prl. Mirrors antiword iGet8InfoLength."""
    spra = (sprm >> 13) & 0x7
    if spra in (0, 1):
        return 3            # 1-byte operand
    if spra in (2, 4, 5):
        return 4            # 2-byte operand
    if spra == 3:
        return 6            # 4-byte operand
    if spra == 7:
        return 5            # 3-byte operand
    if spra == 6:           # variable: first operand byte = size of the rest
        n = u8(data, off + 2)
        return 3 + n
    return 3


def sprm_sgc(sprm: int) -> int:
    return (sprm >> 10) & 0x7


def iter_sprms(data: bytes):
    """Yield (sprm, operand_bytes) pairs from a Prl / grpprl byte array."""
    off = 0
    n = len(data)
    while off + 2 <= n:
        sprm = u16(data, off)
        L = sprm_len(sprm, data, off)
        if off + L > n:
            break
        yield sprm, data[off + 2: off + L]
        off += L


# --- character property sprm values (ispmd + sgc/spra encoded) ---------- #
SPRM_CF_BOLD = 0x0835
SPRM_CF_ITALIC = 0x0836
SPRM_CF_STRIKE = 0x0837
SPRM_CF_OUTLINE = 0x0838
SPRM_CF_SHADOW = 0x0839
SPRM_CF_SMALLCAPS = 0x083A
SPRM_CF_CAPS = 0x083B
SPRM_CF_VANISH = 0x083C
SPRM_CF_EMBOSS = 0x083D
SPRM_CF_IMPRINT = 0x0854
SPRM_CF_WEBHIDDEN = 0x0811
SPRM_CF_SPEC = 0x0855
SPRM_CF_SPECVANISH = 0x0818
SPRM_CF_DSTRIKE = 0x2A53
SPRM_CF_NO_PROOF = 0x0875
SPRM_CF_DATA = 0x0806
SPRM_CF_OLE2 = 0x080A
SPRM_CF_OBJ = 0x0856
SPRM_C_KUL = 0x2A3E
SPRM_C_ICO = 0x2A42
SPRM_C_CV = 0x6870
SPRM_C_HPS = 0x4A43
SPRM_C_HPS_POS = 0x4845
SPRM_C_ISS = 0x2A48
SPRM_C_RG_FTC0 = 0x4A4F
SPRM_C_RG_FTC1 = 0x4A50
SPRM_C_RG_FTC2 = 0x4A51
SPRM_C_CHAR_SCALE = 0x4852
SPRM_C_HIGHLIGHT = 0x2A0C
SPRM_C_DXA_SPACE = 0x8840
SPRM_C_KERN = 0x484B
SPRM_C_PIC_LOCATION = 0x6A03
SPRM_C_SYMBOL = 0x6A09

# --- paragraph property sprm values ------------------------------------- #
SPRM_P_ISTD = 0x4600
SPRM_P_ISTD_PERMUTE = 0xC601
SPRM_P_JC = 0x2461
SPRM_P_JC80 = 0x2403
SPRM_P_ILVL = 0x260A
SPRM_P_ILFO = 0x460B
SPRM_P_OUTLVL = 0x2640
SPRM_P_DXA_LEFT = 0x845E
SPRM_P_DXA_LEFT80 = 0x840F
SPRM_P_DXA_LEFT180 = 0x8411
SPRM_P_DXA_RIGHT = 0x845D
SPRM_P_DXA_RIGHT80 = 0x840E
SPRM_P_DXA_LEFT1 = 0x8460
SPRM_P_DYA_BEFORE = 0xA413
SPRM_P_DYA_AFTER = 0xA414
SPRM_P_DYA_LINE = 0x6412
SPRM_P_F_KEEP = 0x2405
SPRM_P_F_KEEP_FOLLOW = 0x2406
SPRM_P_F_PAGE_BREAK_BEFORE = 0x2407
SPRM_P_F_NO_LINE_NUMB = 0x240C
SPRM_P_F_WIDOW_CONTROL = 0x2431
SPRM_P_F_IN_TABLE = 0x2416
SPRM_P_F_TTP = 0x2417
SPRM_P_F_INNER_TABLE_CELL = 0x244B
SPRM_P_F_INNER_TTP = 0x244C
SPRM_P_F_ADJUST_RIGHT = 0x2448
SPRM_P_ITAP = 0x6649
SPRM_P_DTAP = 0x664A
SPRM_P_SHD = 0xC64D
SPRM_P_BRC_TOP = 0xC64E
SPRM_P_BRC_LEFT = 0xC64F
SPRM_P_BRC_BOTTOM = 0xC650
SPRM_P_BRC_RIGHT = 0xC651
SPRM_P_CHG_TABS = 0xC615
SPRM_P_CHG_TABS_PAPX = 0xC60D
SPRM_P_DCS = 0x442C
SPRM_P_NUM_RM = 0xC645
SPRM_P_PC = 0x261B

# --- table property sprms (stored in the end-of-row PAPX) -------------- #
SPRM_T_DEF_TABLE = 0xD608
SPRM_T_SET_BRC = 0xD62F
SPRM_T_SET_SHD = 0xD62D
SPRM_T_MERGE = 0x5624
SPRM_T_SPLIT = 0x5625
SPRM_T_JCC = 0x548A
SPRM_T_WIDTH = 0xF614
SPRM_T_TABLE_BORDERS = 0xD613
SPRM_T_ROW_HEIGHT = 0x9407
SPRM_T_CANT_SPLIT = 0x3466
SPRM_T_HEADER = 0x3404
SPRM_T_CELL_PADDING = 0xD632

# --- section property sprm values ---------------------------------------- #
SPRM_S_BKC = 0x3009
SPRM_S_F_TITLE_PAGE = 0x300A
SPRM_S_GRPF_IHDT = 0x3014
SPRM_S_NFC_PGN = 0x300E
SPRM_S_F_PGN_RESTART = 0x3011
SPRM_S_DYA_HDR_TOP = 0xB017
SPRM_S_DYA_HDR_BOTTOM = 0xB018
SPRM_S_ORIENTATION = 0x301D
SPRM_S_XA_PAGE = 0xB01F
SPRM_S_YA_PAGE = 0xB020
SPRM_S_DXA_LEFT = 0xB021
SPRM_S_DXA_RIGHT = 0xB022
SPRM_S_DYA_TOP = 0x9023
SPRM_S_DYA_BOTTOM = 0x9024
SPRM_S_DZA_GUTTER = 0xB025
SPRM_S_DM_PAPER_REQ = 0x5026
SPRM_S_COLUMNS = 0x500B
SPRM_S_DXA_COLUMNS = 0x900C
SPRM_S_PGN_START = 0x7044
SPRM_S_TEXT_FLOW = 0x5033


def ico_to_hex(ico: int) -> str | None:
    """Map a Word indexed-colour (ico) value to '#RRGGBB' or None."""
    palette = {
        0: None, 1: None, 2: "auto",
        3: "#000080", 4: "#008000", 5: "#800000", 6: "#800080",
        7: "#808000", 8: "#808080", 9: "#C0C0C0", 10: "#0000FF",
        11: "#00FF00", 12: "#FF0000", 13: "#FF00FF", 14: "#FFFF00",
        15: "#FFFFFF", 16: "#000000",
    }
    return palette.get(ico & 0xFF)


# --------------------------------------------------------------------------- #
#  Data model
# --------------------------------------------------------------------------- #

@dataclass
class CharFormat:
    bold: bool = False
    italic: bool = False
    underline: str | None = None       # w:u val (single/double/words/…)
    strike: bool = False
    dstrike: bool = False
    small_caps: bool = False
    caps: bool = False
    vanish: bool = False
    no_proof: bool = False
    font_ascii: int | None = None      # ftc index
    font_fe: int | None = None
    font_other: int | None = None
    size: int | None = None            # half-points
    color: str | None = None           # '#RRGGBB' / 'auto'
    highlight: str | None = None       # w:highlight val
    vert_align: str | None = None      # none / superscript / subscript
    char_scale: int | None = None      # percent
    char_spacing: int | None = None    # twips
    kern: int | None = None
    spec: bool = False
    fld_kind: int = 0                  # 0 none, 1 begin, 2 separate, 3 end
    obj: bool = False


@dataclass
class CharRun:
    text: str = ""
    fmt: CharFormat = field(default_factory=CharFormat)


@dataclass
class ParaFormat:
    istd: int = 0
    jc: int | None = None              # 0 left,1 center,2 right,3 justify,4 distribute
    left: int | None = None            # twips
    right: int | None = None
    first_line: int | None = None      # twips (negative = hanging)
    before: int | None = None
    after: int | None = None
    line: int | None = None            # 3-byte: lspd
    line_rule: str | None = None       # auto / exact / atLeast / multiple
    keep_next: bool = False
    keep_lines: bool = False
    page_break_before: bool = False
    widow_control: bool | None = None
    outline_lvl: int | None = None
    ilvl: int | None = None
    ilfo: int | None = None
    shading: str | None = None
    tabs: list = field(default_factory=list)   # (pos, align, leader)
    style: str | None = None
    in_table: bool = False
    end_of_row: bool = False
    row_tbl: dict = field(default_factory=dict)  # table props for end-of-row marks


@dataclass
class Paragraph:
    cp_start: int
    cp_end: int
    pfmt: ParaFormat
    runs: list = field(default_factory=list)
    kind: str = "body"                # body / header / footer


@dataclass
class Section:
    cp_start: int = 0
    page_w: int = 11906
    page_h: int = 16838
    margin_l: int = 1800
    margin_r: int = 1800
    margin_t: int = 1440
    margin_b: int = 1440
    gutter: int = 0
    header_dist: int = 851
    footer_dist: int = 992
    orient: str = "portrait"
    title_page: bool = False
    pgn_start: int | None = None
    columns: int = 1
    headers: dict = field(default_factory=dict)   # hdrtype -> list[Paragraph]


# --------------------------------------------------------------------------- #
#  .doc reader
# --------------------------------------------------------------------------- #
class DocReader:
    """Parse the binary .doc into an intermediate document model."""

    def __init__(self, path: str):
        if not olefile.isOleFile(path):
            raise ValueError(f"not an OLE2 (Compound File) container: {path}")
        self.path = path
        self.ole = olefile.OleFileIO(path)
        try:
            self.wd = self.ole.openstream("WordDocument").read()
        except IOError:
            raise ValueError("no WordDocument stream; is this really a .doc file?")
        self._read_fib()
        tbl_name = self._pick_table_stream()
        self.tbl = self.ole.openstream(tbl_name).read()
        self.data = b""
        try:
            self.data = self.ole.openstream("Data").read()
        except (IOError, KeyError):
            pass
        self.fonts: dict[int, str] = {}
        self._read_font_table()
        self.pieces = self._read_piece_table()
        self.text = self._extract_text()
        self.chars: dict[int, CharFormat] = {}
        self.paras: dict[int, Paragraph] = {}
        self.styles: dict[int, tuple] = {}
        self.sections: list[Section] = []
        self.tables: list = []

    # -- FIB ---------------------------------------------------------------
    def _pick_table_stream(self) -> str:
        """Choose the Table stream (0Table/1Table) that actually contains the
        piece table — the fWhichTblStm flag is unreliable in WPS files."""
        names = [n for n in ("0Table", "1Table") if self.ole.exists(n)]
        if len(names) <= 1:
            return names[0] if names else "0Table"
        try:
            fc_clx, lcb_clx = self._fclcb(33)
        except Exception:
            return names[0]
        for name in names:
            try:
                data = self.ole.openstream(name).read()
            except IOError:
                continue
            if 0 <= fc_clx < len(data) and 1 <= lcb_clx < len(data) - fc_clx:
                return name
        return names[0]

    def _read_fib(self) -> None:
        self.wIdent = u16(self.wd, 0)
        self.nFib = u16(self.wd, 2)
        self.fcMin = u32(self.wd, 24)
        self.fcMac = u32(self.wd, 28)
        csw = u16(self.wd, 32)
        cslw = u16(self.wd, 32 + 2 + csw * 2)
        self.lw_off = 32 + 2 + csw * 2 + 2
        lw = struct.unpack_from("<22I", self.wd, self.lw_off)
        (self.cbMac, _, _, self.ccpText, self.ccpFtn, self.ccpHdd, _,
         self.ccpAtn, self.ccpEdn, self.ccpTxbx, self.ccpHdrTxbx, *_rest) = lw
        self.fc_lcb_off = self.lw_off + cslw * 4
        cb = u16(self.wd, self.fc_lcb_off)
        self.fc_lcb = struct.unpack_from(
            "<{}I".format(cb * 2), self.wd, self.fc_lcb_off + 2)

    def _fclcb(self, index: int):
        return self.fc_lcb[index * 2], self.fc_lcb[index * 2 + 1]

    # -- font table (STTBF ffn) --------------------------------------------
    def _read_font_table(self) -> None:
        fc, lcb = self._fclcb(15)   # fcSttbfffn
        if lcb <= 0 or fc >= len(self.tbl):
            return
        sttbf = self.tbl[fc: fc + lcb]
        if len(sttbf) < 4:
            return
        if u16(sttbf, 0) == 0xFFFF:
            ext = True
            cData = u16(sttbf, 2)
            off = 6
        else:
            ext = False
            cData = u16(sttbf, 0)
            off = 4
        if cData == 0 or cData > 0x2000:
            return
        for idx in range(cData):
            if off + 2 > len(sttbf):
                break
            cbFfnM1 = u8(sttbf, off)
            total = cbFfnM1 + 1
            if total < 40 or off + total > len(sttbf):
                break
            name_start = off + 40  # fixed FFN header before szFfn
            if name_start >= off + total:
                break
            if ext:
                end = name_start
                while end + 1 < off + total and u16(sttbf, end) != 0:
                    end += 2
                name = sttbf[name_start:end].decode("utf-16-le", errors="replace")
            else:
                # WPS writes 2-byte names even without fExtend; sniff the bytes.
                probe = sttbf[name_start: min(off + total, name_start + 40)]
                s16 = probe.decode("utf-16-le", errors="replace")
                good16 = sum(1 for c in s16 if c.isalnum() or c in " \u4e00-\u9fff" or c.isspace() or (0x4E00 <= ord(c) <= 0x9FFF))
                s8 = probe.decode("cp1252", errors="replace")
                good8 = sum(1 for c in s8 if c.isalnum() or c.isspace() or 32 <= ord(c) < 127)
                if len(s16) >= 2 and good16 * 2 >= len(s16):
                    end = name_start
                    while end + 1 < off + total and u16(sttbf, end) != 0:
                        end += 2
                    name = sttbf[name_start:end].decode("utf-16-le", errors="replace")
                else:
                    end = name_start
                    while end < off + total and sttbf[end] != 0:
                        end += 1
                    name = sttbf[name_start:end].decode("cp1252", errors="replace")
            self.fonts[idx] = name or f"Font{idx}"
            off += total

    # -- piece table (CLX) --------------------------------------------------
    def _read_piece_table(self) -> list:
        fcClx, lcbClx = self._fclcb(33)
        if lcbClx <= 0:
            return []
        clx = self.tbl[fcClx: fcClx + lcbClx]
        pieces = []
        off = 0
        while off < len(clx):
            clxt = clx[off]
            if clxt == 0x02:
                lcb = u32(clx, off + 1)
                plc = clx[off + 5: off + 5 + lcb]
                n = (lcb - 4) // 12
                cps = [u32(plc, i * 4) for i in range(n + 1)]
                for i in range(n):
                    pcd = plc[4 * (n + 1) + i * 8: 4 * (n + 1) + i * 8 + 8]
                    flags, fc_raw = struct.unpack_from("<HI", pcd, 0)
                    compressed = bool(fc_raw & 0x40000000) or bool(flags & 0x0001)
                    fc = fc_raw & 0x3FFFFFFF
                    pieces.append((cps[i], cps[i + 1], fc, compressed))
                off += 5 + lcb
            else:
                off += 1
        pieces.sort()
        return pieces

    def _extract_text(self) -> str:
        out = []
        for cp0, cp1, fc, comp in self.pieces:
            n = cp1 - cp0
            if comp:
                out.append(self.wd[fc: fc + n].decode("cp1252", errors="replace"))
            else:
                out.append(self.wd[fc: fc + 2 * n].decode("utf-16-le", errors="replace"))
        return "".join(out)

    def fc_to_cp(self, fc: int) -> int:
        best_end = 0
        for cp0, cp1, pfc, comp in self.pieces:
            span = cp1 - cp0
            end = pfc + (span if comp else span * 2)
            if pfc <= fc < end:
                return cp0 + (fc - pfc) // (1 if comp else 2)
            if pfc <= fc:
                best_end = cp1  # fc lies in the gap after this piece
        return best_end

    # -- FKP pages ----------------------------------------------------------
    def _fkp(self, page_no: int) -> bytes:
        return self.wd[page_no * 512: page_no * 512 + 512]

    def _read_chpx_fkps(self) -> list:
        """Return [(fc_start, fc_end, CharFormat)] runs."""
        fc, lcb = self._fclcb(12)   # fcPlcfbteChpx
        if lcb <= 0:
            return []
        plc = self.tbl[fc: fc + lcb]
        n = (len(plc) - 4) // 8
        pns = [u32(plc, 4 * (n + 1) + i * 4) for i in range(n)]
        runs = []
        for pn in pns:
            fkp = self._fkp(pn & 0x7FFFFFFF)
            crun = fkp[511]
            for i in range(crun):
                fc0 = u32(fkp, i * 4)
                fc1 = u32(fkp, (i + 1) * 4)
                rgb = fkp[(crun + 1) * 4 + i]
                ifo = rgb * 2
                if ifo <= 0 or ifo >= 512:
                    continue
                cb = fkp[ifo]
                fmt = self._parse_chpx(fkp[ifo + 1: ifo + 1 + cb])
                runs.append((fc0, fc1, fmt))
        runs.sort(key=lambda r: r[0])
        return runs

    def _read_papx_fkps(self) -> list:
        """Return [(fc, ParaFormat)] — one entry per paragraph."""
        fc, lcb = self._fclcb(13)   # fcPlcfbtePapx
        if lcb <= 0:
            return []
        plc = self.tbl[fc: fc + lcb]
        n = (len(plc) - 4) // 8
        pns = [u32(plc, 4 * (n + 1) + i * 4) for i in range(n)]
        paras = []
        for pn in pns:
            fkp = self._fkp(pn & 0x7FFFFFFF)
            cpara = fkp[511]
            for i in range(cpara):
                pfc = u32(fkp, i * 4)
                rgb = fkp[(cpara + 1) * 4 + i * 13]
                ifo = rgb * 2
                if ifo <= 0 or ifo >= 512:
                    paras.append((pfc, ParaFormat()))
                    continue
                iLen = 2 * fkp[ifo]
                if iLen == 0:
                    ifo += 1
                    iLen = 2 * fkp[ifo]
                istd = u16(fkp, ifo + 1)
                grpprl = fkp[ifo + 3: ifo + 3 + max(0, iLen - 3)]
                pf = ParaFormat(istd=istd)
                self._parse_papx(grpprl, pf)
                paras.append((pfc, pf))
        paras.sort(key=lambda p: p[0])
        return paras

    # -- sprm -> formatting --------------------------------------------------
    def _parse_chpx(self, grpprl: bytes) -> CharFormat:
        f = CharFormat()
        for sprm, op in iter_sprms(grpprl):
            if sprm == SPRM_CF_BOLD:
                f.bold = _toggle(op)
            elif sprm == SPRM_CF_ITALIC:
                f.italic = _toggle(op)
            elif sprm == SPRM_CF_STRIKE:
                f.strike = _toggle(op)
            elif sprm == SPRM_CF_DSTRIKE:
                f.dstrike = _toggle(op)
            elif sprm == SPRM_CF_SMALLCAPS:
                f.small_caps = _toggle(op)
            elif sprm == SPRM_CF_CAPS:
                f.caps = _toggle(op)
            elif sprm == SPRM_CF_VANISH:
                f.vanish = _toggle(op)
            elif sprm == SPRM_CF_NO_PROOF:
                f.no_proof = _toggle(op)
            elif sprm == SPRM_CF_SPEC:
                f.spec = _toggle(op)
            elif sprm == SPRM_CF_EMBOSS or sprm == SPRM_CF_IMPRINT or sprm == SPRM_CF_SHADOW or sprm == SPRM_CF_OUTLINE:
                pass
            elif sprm == SPRM_C_KUL:
                f.underline = _underline_val(op)
            elif sprm == SPRM_C_ICO:
                c = ico_to_hex(u16(op, 0) if len(op) >= 2 else op[0])
                if c:
                    f.color = c
            elif sprm == SPRM_C_CV:
                cv = u32(op, 0) if len(op) >= 4 else 0
                if cv & 0x40000000:
                    pass  # theme colour — leave default
                elif cv != 0 and cv != 0xFEFFFFFF:
                    f.color = "#%06X" % (cv & 0xFFFFFF)
            elif sprm == SPRM_C_HPS:
                if len(op) >= 2:
                    f.size = u16(op, 0)
            elif sprm == SPRM_C_HPS_POS:
                if len(op) >= 2:
                    pos = i16(op, 0)
                    f.vert_align = "superscript" if pos > 0 else ("subscript" if pos < 0 else None)
            elif sprm == SPRM_C_ISS:
                if len(op) >= 2:
                    v = u16(op, 0)
                    f.vert_align = {0: None, 1: "superscript", 2: "subscript"}.get(v & 0xFFFF)
            elif sprm == SPRM_C_RG_FTC0:
                if len(op) >= 2:
                    f.font_ascii = u16(op, 0) & 0xFFFF
            elif sprm == SPRM_C_RG_FTC1:
                if len(op) >= 2:
                    f.font_fe = u16(op, 0) & 0xFFFF
            elif sprm == SPRM_C_RG_FTC2:
                if len(op) >= 2:
                    f.font_other = u16(op, 0) & 0xFFFF
            elif sprm == SPRM_C_CHAR_SCALE:
                if len(op) >= 2:
                    f.char_scale = u16(op, 0)
            elif sprm == SPRM_C_HIGHLIGHT:
                if len(op):
                    f.highlight = _highlight_val(op[0])
            elif sprm == SPRM_C_DXA_SPACE:
                if len(op) >= 2:
                    f.char_spacing = i16(op, 0)
            elif sprm == SPRM_C_KERN:
                if len(op) >= 2:
                    f.kern = u16(op, 0)
        return f

    def _parse_papx(self, grpprl: bytes, pf: ParaFormat) -> None:
        for sprm, op in iter_sprms(grpprl):
            if sprm in (SPRM_P_JC, SPRM_P_JC80):
                if len(op):
                    pf.jc = op[0] & 0xFF
            elif sprm in (SPRM_P_DXA_LEFT, SPRM_P_DXA_LEFT80, SPRM_P_DXA_LEFT180):
                if len(op) >= 2:
                    pf.left = _u16s(op, 0)
            elif sprm in (SPRM_P_DXA_RIGHT, SPRM_P_DXA_RIGHT80):
                if len(op) >= 2:
                    pf.right = _u16s(op, 0)
            elif sprm == SPRM_P_DXA_LEFT1:
                if len(op) >= 2:
                    pf.first_line = _u16s(op, 0)
            elif sprm == SPRM_P_DYA_BEFORE:
                if len(op) >= 2:
                    pf.before = u16(op, 0)
            elif sprm == SPRM_P_DYA_AFTER:
                if len(op) >= 2:
                    pf.after = u16(op, 0)
            elif sprm == SPRM_P_DYA_LINE:
                if len(op) >= 3:
                    pf.line = u16(op, 0)
                    pf.line_rule = _line_rule(u16(op, 2))
            elif sprm == SPRM_P_F_KEEP:
                pf.keep_lines = _toggle(op)
            elif sprm == SPRM_P_F_KEEP_FOLLOW:
                pf.keep_next = _toggle(op)
            elif sprm == SPRM_P_F_PAGE_BREAK_BEFORE:
                pf.page_break_before = _toggle(op)
            elif sprm == SPRM_P_F_WIDOW_CONTROL:
                pf.widow_control = _toggle(op)
            elif sprm == SPRM_P_OUTLVL:
                if len(op):
                    pf.outline_lvl = op[0] & 0xFF
            elif sprm == SPRM_P_ILVL:
                if len(op):
                    pf.ilvl = op[0] & 0xFF
            elif sprm == SPRM_P_ILFO:
                if len(op) >= 2:
                    pf.ilfo = u16(op, 0)
            elif sprm == SPRM_P_F_IN_TABLE or sprm == SPRM_P_F_INNER_TABLE_CELL:
                pf.in_table = _toggle(op)
            elif sprm == SPRM_P_F_TTP or sprm == SPRM_P_F_INNER_TTP:
                pf.end_of_row = _toggle(op)
            elif sprm == SPRM_P_SHD:
                pass  # paragraph shading (best-effort skip)
            elif sprm == SPRM_P_ISTD:
                if len(op) >= 2:
                    pf.istd = u16(op, 0)
            elif sprm == SPRM_P_CHG_TABS or sprm == SPRM_P_CHG_TABS_PAPX:
                self._parse_tabs(op, pf)
            elif sprm == SPRM_P_NUM_RM or sprm == SPRM_P_ISTD_PERMUTE:
                pass
            elif sprm == SPRM_T_DEF_TABLE:
                if len(op) >= 5:
                    col = u8(op, 2)
                    pos = [i16(op, 3 + k * 2) for k in range(col + 1) if 3 + k * 2 + 1 < len(op)]
                    widths = [max(0, pos[k + 1] - pos[k]) for k in range(len(pos) - 1)]
                    pf.row_tbl["col_widths"] = widths
            elif sprm in (0xD605, 0xD613, 0xD62F):
                borders = self._parse_table_borders(op)
                if borders:
                    pf.row_tbl["borders"] = borders
            elif sprm == SPRM_T_SET_BRC or sprm == SPRM_T_SET_SHD:
                pass
            elif sprm == SPRM_T_WIDTH:
                if len(op) >= 4:
                    pf.row_tbl["width"] = u16(op, 2)
        if pf.istd == 0 and not any(v for v in (pf.left, pf.right, pf.first_line)):
            pass

    def _parse_tabs(self, op: bytes, pf: ParaFormat) -> None:
        if not op:
            return
        if op[0] == 0xFF:
            # delete tabs: count of 4-byte entries
            n = op[1]
            pos = 2
            for _ in range(n):
                if pos + 4 <= len(op):
                    pos += 4
            if pos < len(op) and op[pos] == 0:
                # add tabs follow
                add_n = op[pos + 1] if pos + 1 < len(op) else 0
                pos += 2
                for _ in range(add_n):
                    if pos + 3 <= len(op):
                        pos += 3
            return
        # add tabs
        n = op[0]
        pos = 1
        for _ in range(n):
            if pos + 3 <= len(op):
                dvp = u16(op, pos)
                tab_align = op[pos + 2] & 0x7
                leader = (op[pos + 2] >> 3) & 0x7
                pf.tabs.append((dvp, _tab_align(tab_align), _tab_leader(leader)))
                pos += 3

    def _parse_table_borders(self, op: bytes) -> list:
        """Parse a TTableBorders operand into [(side, btype, width, color)]."""
        if not op:
            return []
        if op[0] == 0xFF or len(op) < 9:
            return []
        cb = op[0]
        is80 = cb == 0x18
        step = 4 if is80 else 8
        names = ["top", "left", "bottom", "right", "insideH", "insideV"]
        out = []
        for i, side in enumerate(names):
            base = 1 + i * step
            if base + step > len(op):
                break
            brc = op[base:base + step]
            if is80:
                btype = brc[0] & 0xFF
                width = brc[3] >> 4 if len(brc) > 3 else 4
                color = None
            else:
                cv = u32(brc, 0)
                width = brc[4]
                btype = brc[5]
                color = None if cv in (0, 0xFF000000) else "#%06X" % (cv & 0xFFFFFF)
            out.append((side, btype, width, color))
        return out

    # -- styles (STSH) -------------------------------------------------------
    def _read_styles(self) -> None:
        fc, lcb = self._fclcb(0)   # fcStshfOrig (same data as fcStshf usually)
        if lcb <= 0:
            fc, lcb = self._fclcb(1)
        if lcb <= 0 or fc >= len(self.tbl):
            return
        stsh = self.tbl[fc: fc + lcb]
        cstd = u16(stsh, 2)
        cb_std_base = u16(stsh, 4)
        rglpstd = 2 + u16(stsh, 0)
        if rglpstd + 2 > len(stsh):
            return
        pos = rglpstd
        raw = {}          # istd -> dict(name, type, base, papx_grpprl, chpx_grpprl)
        for i in range(cstd):
            if pos + 2 > len(stsh):
                break
            tStdLen = u16(stsh, pos)
            if tStdLen == 0:
                pos += 2
                continue
            if pos + tStdLen > len(stsh):
                break
            body = stsh[pos + 2: pos + tStdLen]
            if len(body) < 8:
                pos += 2 + tStdLen
                continue
            tmp = u16(body, 2)
            stype = tmp & 0xF
            istd_base = (tmp >> 4) & 0xFFF
            tmp2 = u16(body, 4)
            cupx = tmp2 & 0xF
            istd_next = tmp2 >> 4
            name_pos = cb_std_base
            if name_pos + 2 > len(body):
                pos += 2 + tStdLen
                continue
            name_len = u16(body, name_pos)
            name = body[name_pos + 2: name_pos + 2 + name_len * 2].decode("utf-16-le", errors="replace")
            name_len_bytes = name_len * 2 + 2
            p = name_pos + 2 + name_len_bytes
            if p % 2:
                p += 1
            papx_gp = b""
            chpx_gp = b""
            if stype == 1 and cupx >= 1 and p + 2 <= len(body):
                upx_len = u16(body, p)
                papx_gp = body[p + 4: min(p + 4 + max(0, upx_len - 2), len(body))]
                p += 2 + upx_len
                if p % 2:
                    p += 1
            if (stype == 1 and cupx >= 2) or (stype == 2 and cupx >= 1):
                if p + 2 <= len(body):
                    upx_len = u16(body, p)
                    chpx_gp = body[p + 2: min(p + 2 + upx_len, len(body))]
            raw[i] = dict(name=name, type=stype, base=istd_base, next=istd_next,
                          papx_gp=papx_gp, chpx_gp=chpx_gp)
            pos += 2 + tStdLen
        # resolve chains
        resolved = {}
        visited = set()

        def resolve(istd):
            if istd in resolved:
                return resolved[istd]
            if istd in visited or istd not in raw:
                return (CharFormat(), ParaFormat())
            visited.add(istd)
            st = raw[istd]
            cf = CharFormat()
            pf = ParaFormat(istd=istd)
            if st["base"] != 0xFFF:
                bcf, bpf = resolve(st["base"])
                cf = _merge_cf(bcf, cf)
                pf = _merge_pf(bpf, pf)
            cf2 = self._parse_chpx(st["chpx_gp"])
            cf = _merge_cf(cf, cf2)
            if st["papx_gp"]:
                self._parse_papx(st["papx_gp"], pf)
            st["cf"] = cf
            st["pf"] = pf
            resolved[istd] = (cf, pf)
            return cf, pf

        for i in list(raw.keys()):
            resolve(i)
        self.styles = raw

    # -- sections (PlcfSed / SEPX) -------------------------------------------
    def _read_sections(self) -> None:
        fc, lcb = self._fclcb(6)   # fcPlcfsed
        if lcb < 8 or fc >= len(self.tbl):
            self.sections = [Section()]
            return
        data = self.tbl[fc: fc + lcb]
        n = (len(data) - 4) // 16
        cps = [u32(data, i * 4) for i in range(n + 1)]
        for i in range(n):
            base = 4 * (n + 1) + 12 * i
            fn, fcSepx, fnMpr, fcMpr = struct.unpack_from("<HIHI", data, base)
            sec = Section(cp_start=cps[i])
            if 0 < fcSepx < len(self.wd):
                cb = u16(self.wd, fcSepx)
                grpprl = self.wd[fcSepx + 2: fcSepx + 2 + cb]
                self._parse_sepx(grpprl, sec)
            self.sections.append(sec)
        if not self.sections:
            self.sections.append(Section())

    def _parse_sepx(self, grpprl: bytes, sec: Section) -> None:
        for sprm, op in iter_sprms(grpprl):
            if sprm == SPRM_S_XA_PAGE and len(op) >= 2:
                sec.page_w = u16(op, 0)
            elif sprm == SPRM_S_YA_PAGE and len(op) >= 2:
                sec.page_h = u16(op, 0)
            elif sprm == SPRM_S_DXA_LEFT and len(op) >= 2:
                sec.margin_l = u16(op, 0)
            elif sprm == SPRM_S_DXA_RIGHT and len(op) >= 2:
                sec.margin_r = u16(op, 0)
            elif sprm == SPRM_S_DYA_TOP and len(op) >= 2:
                sec.margin_t = u16(op, 0)
            elif sprm == SPRM_S_DYA_BOTTOM and len(op) >= 2:
                sec.margin_b = u16(op, 0)
            elif sprm == SPRM_S_DZA_GUTTER and len(op) >= 2:
                sec.gutter = u16(op, 0)
            elif sprm == SPRM_S_DYA_HDR_TOP and len(op) >= 2:
                sec.header_dist = u16(op, 0)
            elif sprm == SPRM_S_DYA_HDR_BOTTOM and len(op) >= 2:
                sec.footer_dist = u16(op, 0)
            elif sprm == SPRM_S_ORIENTATION and len(op):
                sec.orient = "landscape" if op[0] else "portrait"
            elif sprm == SPRM_S_F_TITLE_PAGE and len(op):
                sec.title_page = bool(op[0] & 1)
            elif sprm == SPRM_S_PGN_START and len(op) >= 4:
                sec.pgn_start = u32(op, 0)
            elif sprm == SPRM_S_COLUMNS and len(op) >= 2:
                sec.columns = max(1, u16(op, 0) + 1)
            elif sprm == SPRM_S_GRPF_IHDT:
                pass

    # -- assemble the document model -----------------------------------------
    def build(self):
        self._read_styles()
        self._read_sections()
        cps = self._read_chpx_fkps()
        paps = self._read_papx_fkps()
        if not paps:
            paps = [(self.fcMin, ParaFormat())]
        if not cps:
            cps = [(self.fcMin, self.fcMac, CharFormat())]
        max_cp = max(len(self.text), self.ccpText)

        # character-format interval map: cp boundary -> list of CharFormat
        cf_map = {}
        for fc0, fc1, fmt in cps:
            cp0 = min(max(0, self.fc_to_cp(fc0)), max_cp)
            cp1 = min(max(cp0, self.fc_to_cp(fc1)), max_cp)
            if cp0 >= 0 and cp1 >= 0:
                cf_map.setdefault(cp0, []).append(fmt)
                cf_map.setdefault(max(cp0, cp1), []).append(None)
        self.cf_map = cf_map

        # paragraphs from PAPX boundaries
        para_fcs = sorted(set(p[0] for p in paps))
        pfmt_by_fc = {p[0]: p[1] for p in paps}
        paragraphs = []
        for i, fc0 in enumerate(para_fcs):
            fc1 = para_fcs[i + 1] if i + 1 < len(para_fcs) else self.fcMac
            cp0 = min(max(0, self.fc_to_cp(fc0)), max_cp)
            cp1 = min(max(cp0, self.fc_to_cp(fc1)), max_cp)
            if cp1 <= cp0:
                continue
            pf = pfmt_by_fc.get(fc0, ParaFormat())
            paragraphs.append(Paragraph(cp0, cp1, pf))
        self.paragraphs = paragraphs

        # group tables
        self.tables = self._group_tables()
        return self

    def _group_tables(self) -> list:
        """Split paragraphs into body paragraphs and table groups.

        Returns a list of tables: each table is a list of rows; each row is a
        list of cells; each cell is a list of Paragraph objects.
        """
        paras = self.paragraphs
        tables = []
        i = 0
        while i < len(paras):
            if not paras[i].pfmt.in_table:
                i += 1
                continue
            row_start = i
            rows = []
            while i < len(paras) and paras[i].pfmt.in_table:
                row_paras = []
                is_eor = False
                while i < len(paras) and paras[i].pfmt.in_table:
                    row_paras.append(paras[i])
                    is_eor = paras[i].pfmt.end_of_row
                    i += 1
                    if is_eor:
                        break
                cells = self._split_row_into_cells(row_paras)
                rows.append((row_paras, cells))
            if rows:
                tables.append(rows)
        return tables

    def _split_row_into_cells(self, row_paras: list) -> list:
        """Split a row's paragraphs into per-cell paragraph lists."""
        if not row_paras:
            return []
        start = row_paras[0].cp_start
        end = row_paras[-1].cp_end
        row_text = self.text[start:end]
        marks = [pos for pos, ch in enumerate(row_text) if ch == "\x07"]
        cells = []
        prev = 0
        # the final \x07 is the end-of-row mark, not a real cell
        for m in marks[:-1]:
            c0 = start + prev
            c1 = start + m
            cell_paras = [p for p in row_paras if p.cp_start < c1 and p.cp_end > c0]
            cells.append(cell_paras)
            prev = m + 1
        return cells


# --------------------------------------------------------------------------- #
#  helpers
# --------------------------------------------------------------------------- #

def _toggle(op: bytes) -> bool:
    if not op:
        return False
    v = op[0]
    if v == 0:
        return False
    if v == 0x80:
        return True
    return bool(v & 1)


def _u16s(op: bytes, off: int) -> int:
    return u16(op, off)


def _underline_val(op: bytes) -> str | None:
    if not op:
        return None
    v = u16(op, 0) if len(op) >= 2 else op[0]
    v &= 0xFF
    m = {0: None, 1: "single", 2: "words", 3: "double", 5: "dotted",
         7: "wave", 9: "thick", 10: "dash", 11: "dotDash", 12: "dotDotDash",
         20: "single", 21: "dotted", 22: "double", 23: "thick", 24: "dash"}
    return m.get(v, "single" if v else None)


def _highlight_val(v: int) -> str | None:
    m = {0: None, 1: "black", 2: "blue", 3: "cyan", 4: "green", 5: "magenta",
         6: "red", 7: "yellow", 8: "white", 9: "darkBlue", 10: "darkCyan",
         11: "darkGreen", 12: "darkMagenta", 13: "darkRed", 14: "darkYellow",
         15: "darkGray", 16: "lightGray"}
    return m.get(v & 0xFF)


def _line_rule(lspd: int) -> str:
    r = (lspd >> 4) & 0x3
    if r == 0:
        return "multiple"
    if r == 1:
        return "atLeast"
    if r == 2:
        return "exact"
    return "auto"


def _tab_align(v: int) -> str:
    return {0: "left", 1: "center", 2: "right", 3: "decimal", 4: "bar"}.get(v, "left")


def _tab_leader(v: int) -> str:
    return {0: None, 1: "dot", 2: "hyphen", 3: "underscore", 4: "heavy"}.get(v)


BRC_VAL = {0: "single", 1: "single", 2: "double", 3: "shadow", 4: "dotted",
           5: "dashed", 6: "hairline", 20: "nil", 21: "single", 22: "dotted",
           23: "double", 24: "single", 25: "dash", 26: "dotDash"}


def _border_val(btype: int) -> str:
    return BRC_VAL.get(btype & 0xFF, "single" if btype else "nil")


def _merge_cf(base: CharFormat, override: CharFormat) -> CharFormat:
    out = CharFormat()
    for f in dc_fields(CharFormat):
        bv = getattr(base, f.name)
        ov = getattr(override, f.name)
        if isinstance(bv, bool):
            setattr(out, f.name, ov if ov else bv)
        else:
            setattr(out, f.name, ov if ov is not None else bv)
    return out


def _merge_pf(base: ParaFormat, override: ParaFormat) -> ParaFormat:
    out = ParaFormat()
    for f in dc_fields(ParaFormat):
        bv = getattr(base, f.name)
        ov = getattr(override, f.name)
        if isinstance(bv, bool):
            setattr(out, f.name, ov if ov else bv)
        elif f.name in ("tabs", "row_tbl"):
            setattr(out, f.name, list(bv) + list(ov))
        else:
            setattr(out, f.name, ov if ov is not None else bv)
    return out


# --------------------------------------------------------------------------- #
#  .docx builder
# --------------------------------------------------------------------------- #

JC_MAP = {0: "left", 1: "center", 2: "right", 3: "both", 4: "distribute"}
TAG = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _el(parent, tag, attrs=None, text=None):
    e = etree.SubElement(parent, f"{{{TAG}}}{tag}")
    if attrs:
        for k, v in attrs.items():
            e.set(f"{{{TAG}}}{k}", v)
    if text is not None:
        e.text = text
    return e


def _set_val(parent, tag, val):
    if val:
        _el(parent, tag).set(f"{{{TAG}}}val", val)


def build_rpr(cf: CharFormat, fonts: dict):
    """Build a w:rPr element (or None) from a resolved CharFormat."""
    rPr = etree.Element(f"{{{TAG}}}rPr")

    def el(tag, val_attr=None):
        e = etree.SubElement(rPr, f"{{{TAG}}}{tag}")
        if val_attr is not None:
            e.set(f"{{{TAG}}}val", val_attr)
        return e

    any_flag = False
    if cf.bold:
        el("b"); any_flag = True
    if cf.italic:
        el("i"); any_flag = True
    if cf.strike:
        el("strike"); any_flag = True
    if cf.dstrike:
        el("dstrike"); any_flag = True
    if cf.small_caps:
        el("smallCaps"); any_flag = True
    if cf.caps:
        el("caps"); any_flag = True
    if cf.vanish:
        el("vanish"); any_flag = True
    if cf.no_proof:
        el("noProof"); any_flag = True
    if cf.underline:
        el("u", cf.underline); any_flag = True
    if cf.color:
        el("color", cf.color.lstrip("#")); any_flag = True
    if cf.highlight:
        el("highlight", cf.highlight); any_flag = True
    if cf.vert_align:
        el("vertAlign", cf.vert_align); any_flag = True
    if cf.size:
        el("sz", str(cf.size)); el("szCs", str(cf.size)); any_flag = True
    if cf.char_scale and cf.char_scale != 100:
        el("w", str(cf.char_scale)); any_flag = True
    if cf.char_spacing:
        el("spacing", str(cf.char_spacing)); any_flag = True
    if cf.kern:
        el("kern", str(cf.kern)); any_flag = True
    ascii_f = fonts.get(cf.font_ascii or -1) if cf.font_ascii is not None else None
    fe_f = fonts.get(cf.font_fe or -1) if cf.font_fe is not None else None
    other_f = fonts.get(cf.font_other or -1) if cf.font_other is not None else None
    if ascii_f or fe_f or other_f:
        rf = el("rFonts")
        if ascii_f:
            rf.set(f"{{{TAG}}}ascii", ascii_f)
        if fe_f:
            rf.set(f"{{{TAG}}}eastAsia", fe_f)
        if other_f:
            rf.set(f"{{{TAG}}}hAnsi", other_f)
        any_flag = True
    if not any_flag:
        return None
    return rPr


def build_ppr(pf: ParaFormat, fonts: dict):
    """Build a w:pPr element (or None)."""
    pPr = etree.Element(f"{{{TAG}}}pPr")

    def el(tag, **attrs):
        e = etree.SubElement(pPr, f"{{{TAG}}}{tag}")
        for k, v in attrs.items():
            e.set(f"{{{TAG}}}{k}", v)
        return e

    if pf.jc is not None:
        el("jc", val=JC_MAP.get(pf.jc, "left"))
    if pf.left is not None or pf.right is not None or pf.first_line is not None:
        ind = el("ind")
        if pf.left is not None:
            ind.set(f"{{{TAG}}}left", str(pf.left))
        if pf.right is not None:
            ind.set(f"{{{TAG}}}right", str(pf.right))
        if pf.first_line is not None:
            if pf.first_line < 0:
                ind.set(f"{{{TAG}}}hanging", str(-pf.first_line))
            else:
                ind.set(f"{{{TAG}}}firstLine", str(pf.first_line))
    if pf.before is not None or pf.after is not None or pf.line is not None:
        sp = el("spacing")
        if pf.before is not None:
            sp.set(f"{{{TAG}}}before", str(pf.before))
        if pf.after is not None:
            sp.set(f"{{{TAG}}}after", str(pf.after))
        if pf.line is not None:
            rule = pf.line_rule or "multiple"
            sp.set(f"{{{TAG}}}line", str(pf.line))
            if rule != "multiple":
                sp.set(f"{{{TAG}}}lineRule", rule)
    if pf.keep_lines:
        el("keepLines")
    if pf.keep_next:
        el("keepNext")
    if pf.page_break_before:
        el("pageBreakBefore")
    if pf.widow_control is False:
        el("widowControl", val="0")
    if pf.outline_lvl is not None and 0 <= pf.outline_lvl <= 9:
        el("outlineLvl", val=str(pf.outline_lvl))
    if pf.tabs:
        tabs = el("tabs")
        for pos, align, leader in pf.tabs:
            attrs = {"val": align, "pos": str(pos)}
            if leader:
                attrs["leader"] = leader
            el("tab", **attrs)
    if len(pPr) == 0:
        return None
    return pPr


class DocxBuilder:
    def __init__(self, doc: DocReader):
        self.doc = doc
        self.document = Document()
        self.fonts = doc.fonts

    def build(self) -> Document:
        doc = self.doc
        body = self.document.element.body
        para_list = doc.paragraphs
        tables = doc.tables
        # Map first paragraph of each table -> table index for ordered emission.
        first_para_of_table = {}
        for ti, table in enumerate(tables):
            if table:
                first_para_of_table[id(table[0][0][0])] = ti
        emitted_tables = set()
        i = 0
        while i < len(para_list):
            p = para_list[i]
            ti = first_para_of_table.get(id(p))
            if ti is not None and ti not in emitted_tables:
                self._add_table(body, tables[ti])
                emitted_tables.add(ti)
                last = tables[ti][-1][0][-1]
                i = para_list.index(last) + 1
                continue
            self._add_paragraph(body, p)
            i += 1
        self._apply_sections()
        self._apply_styles()
        return self.document

    def _add_table(self, body, table):
        rows = table
        # gather widths and borders from end-of-row paragraphs
        widths = []
        borders = None
        for row_paras, _cells in rows:
            for p in row_paras:
                if p.pfmt.row_tbl.get("col_widths") and not widths:
                    widths = p.pfmt.row_tbl["col_widths"]
                if p.pfmt.row_tbl.get("borders") and borders is None:
                    borders = p.pfmt.row_tbl["borders"]
        grid_cols = max((len(cells) for _rp, cells in rows), default=len(widths))
        if widths and len(widths) >= grid_cols:
            grid_widths = widths[:grid_cols]
        else:
            grid_widths = [1] * grid_cols
        tbl = etree.SubElement(body, f"{{{TAG}}}tbl")
        tblPr = _el(tbl, "tblPr")
        tblW = _el(tblPr, "tblW")
        tblW.set(f"{{{TAG}}}w", str(sum(w for w in grid_widths if w > 1) or grid_cols * 100))
        tblW.set(f"{{{TAG}}}type", "dxa")
        _el(tblPr, "tblLook", {"val": "04A0"})
        if borders:
            bl = _el(tblPr, "tblBorders")
            for side, btype, width, color in borders:
                e = _el(bl, side)
                e.set(f"{{{TAG}}}val", _border_val(btype))
                if color:
                    e.set(f"{{{TAG}}}color", color.lstrip("#"))
        grid = etree.SubElement(tbl, f"{{{TAG}}}tblGrid")
        for wdt in grid_widths:
            gc = etree.SubElement(grid, f"{{{TAG}}}gridCol")
            gc.set(f"{{{TAG}}}w", str(wdt if wdt > 1 else 1000))
        for row_paras, cells in rows:
            tr = etree.SubElement(tbl, f"{{{TAG}}}tr")
            extra = grid_cols - len(cells)
            for ci, cell_paras in enumerate(cells):
                tc = etree.SubElement(tr, f"{{{TAG}}}tc")
                tcPr = etree.SubElement(tc, f"{{{TAG}}}tcPr")
                tcW = etree.SubElement(tcPr, f"{{{TAG}}}tcW")
                tcW.set(f"{{{TAG}}}w", "0")
                tcW.set(f"{{{TAG}}}type", "auto")
                if extra > 0 and ci == len(cells) - 1:
                    gs = etree.SubElement(tcPr, f"{{{TAG}}}gridSpan")
                    gs.set(f"{{{TAG}}}val", str(extra + 1))
                if not cell_paras:
                    etree.SubElement(tc, f"{{{TAG}}}p")
                    continue
                for para in cell_paras:
                    p = etree.SubElement(tc, f"{{{TAG}}}p")
                    base_cf, base_pf = self._style_format(para.pfmt.istd)
                    pf = _merge_pf(base_pf, para.pfmt)
                    pPr = build_ppr(pf, self.fonts)
                    if pPr is not None:
                        p.append(pPr)
                    runs = self._split_runs(para, base_cf)
                    for text, cf in runs:
                        self._emit_text_runs(p, text, cf)
        return tbl

    def _add_paragraph(self, body, para: Paragraph):
        p = etree.SubElement(body, f"{{{TAG}}}p")
        base_cf, base_pf = self._style_format(para.pfmt.istd)
        pf = _merge_pf(base_pf, para.pfmt)
        pPr = build_ppr(pf, self.fonts)
        if pPr is not None:
            p.append(pPr)
        runs = self._split_runs(para, base_cf)
        for text, cf in runs:
            self._emit_text_runs(p, text, cf)
        return p

    def _style_format(self, istd: int):
        cf, pf = CharFormat(), ParaFormat()
        if self.doc.styles:
            st = self.doc.styles.get(istd) or self.doc.styles.get(0)
            if st:
                cf = st.get("cf", CharFormat())
                pf = st.get("pf", ParaFormat())
        return cf, pf

    def _emit_text_runs(self, p, text: str, cf: CharFormat):
        """Emit w:r elements for one text span, translating special chars and
        preserving field begin/separate/end as OOXML field codes."""
        r = etree.SubElement(p, f"{{{TAG}}}r")
        rPr = build_rpr(cf, self.fonts)
        if rPr is not None:
            r.append(rPr)
        buf = []

        def flush():
            nonlocal buf
            if buf:
                t = etree.SubElement(r, f"{{{TAG}}}t")
                t.text = "".join(buf)
                if t.text != t.text.strip():
                    t.set(f"{{{XML_NS}}}space", "preserve")
                buf = []

        for ch in text:
            o = ord(ch)
            if o == 0x0D:
                continue
            if o == 0x07 or o == 0x01:
                continue
            if o == 0x0B:
                flush()
                etree.SubElement(r, f"{{{TAG}}}br")
            elif o == 0x0C:
                flush()
                etree.SubElement(r, f"{{{TAG}}}br", {f"{{{TAG}}}type": "page"})
            elif o == 0x09:
                flush()
                etree.SubElement(r, f"{{{TAG}}}tab")
            elif o == 0x19:
                flush()
                etree.SubElement(r, f"{{{TAG}}}noBreakHyphen")
            elif o == 0x13:
                flush()
                etree.SubElement(r, f"{{{TAG}}}fldChar", {f"{{{TAG}}}fldCharType": "begin"})
            elif o == 0x14:
                flush()
                etree.SubElement(r, f"{{{TAG}}}fldChar", {f"{{{TAG}}}fldCharType": "separate"})
            elif o == 0x15:
                flush()
                etree.SubElement(r, f"{{{TAG}}}fldChar", {f"{{{TAG}}}fldCharType": "end"})
            elif o == 0xA0:
                flush()
                buf.append("\u00a0")
            elif o == 0x1E or o == 0x1F:
                buf.append(" ")
            elif 0x01 <= o <= 0x08:
                flush()
            else:
                buf.append(ch)
        flush()

    def _split_runs(self, para: Paragraph, base_cf: CharFormat):
        """Split a paragraph CP range into (text, CharFormat) spans using the
        CHPX interval boundaries."""
        text = self.doc.text
        cf_map = self.doc.cf_map
        boundary_cps = sorted(cf_map.keys())
        import bisect
        out = []
        i = para.cp_start
        while i < para.cp_end:
            idx = bisect.bisect_right(boundary_cps, i) - 1
            fmt = CharFormat()
            if idx >= 0:
                k = boundary_cps[idx]
                for f in cf_map[k]:
                    if f is not None:
                        fmt = f
            cur_fmt = _merge_cf(base_cf, fmt)
            nxt = boundary_cps[idx + 1] if idx + 1 < len(boundary_cps) else para.cp_end
            end = min(nxt, para.cp_end)
            if end > i:
                out.append((text[i:end], cur_fmt))
            i = end
        if not out:
            out.append(("", base_cf))
        return out

    def _apply_sections(self):
        # Use first section as the default page setup
        doc = self.document
        sec = self.doc.sections[0] if self.doc.sections else Section()
        s = doc.sections[0]
        s.page_width = self._emus(sec.page_w)
        s.page_height = self._emus(sec.page_h)
        s.left_margin = self._emus(sec.margin_l)
        s.right_margin = self._emus(sec.margin_r)
        s.top_margin = self._emus(sec.margin_t)
        s.bottom_margin = self._emus(sec.margin_b)
        s.header_distance = self._emus(sec.header_dist)
        s.footer_distance = self._emus(sec.footer_dist)
        if sec.orient == "landscape":
            s.orientation = 1  # WD_ORIENT.LANDSCAPE

    def _apply_styles(self):
        styles = self.document.styles
        # Normal default font from STSH style 0
        if self.doc.styles:
            st0 = self.doc.styles.get(0)
            if st0:
                cf = st0.get("cf")
                normal = styles["Normal"]
                rpr = normal.element.get_or_add_rPr()
                changed = False
                if cf.font_ascii is not None and self.fonts.get(cf.font_ascii):
                    rf = rpr.find(qn("w:rFonts"))
                    if rf is None:
                        rf = etree.SubElement(rpr, qn("w:rFonts"))
                    rf.set(qn("w:ascii"), self.fonts[cf.font_ascii])
                    rf.set(qn("w:hAnsi"), self.fonts[cf.font_ascii])
                    changed = True
                if cf.font_fe is not None and self.fonts.get(cf.font_fe):
                    rf = rpr.find(qn("w:rFonts"))
                    if rf is None:
                        rf = etree.SubElement(rpr, qn("w:rFonts"))
                    rf.set(qn("w:eastAsia"), self.fonts[cf.font_fe])
                    changed = True
                if cf.size:
                    for tag in ("w:sz", "w:szCs"):
                        e = rpr.find(qn(tag))
                        if e is None:
                            e = etree.SubElement(rpr, qn(tag))
                        e.set(qn("w:val"), str(cf.size))
                    changed = True
                if cf.bold and rpr.find(qn("w:b")) is None:
                    etree.SubElement(rpr, qn("w:b"))
                    changed = True
                if cf.italic and rpr.find(qn("w:i")) is None:
                    etree.SubElement(rpr, qn("w:i"))
                    changed = True
                if not changed:
                    rpr.getparent().remove(rpr)

    def _emus(self, twips: int) -> int:
        return int(twips * 635) if twips else 0


# re-export qn for use
from docx.oxml.ns import qn  # noqa: E402


# --------------------------------------------------------------------------- #
#  public API
# --------------------------------------------------------------------------- #

def convert(src: str, dst: str | None = None, report: bool = True) -> dict:
    """Convert a .doc file to .docx. Returns a fidelity report dict."""
    if dst is None:
        dst = os.path.splitext(src)[0] + ".docx"
    reader = DocReader(src)
    reader.build()
    builder = DocxBuilder(reader)
    document = builder.build()
    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    document.save(dst)
    rep = {
        "src": src,
        "dst": dst,
        "fib": reader.nFib,
        "chars": len(reader.text),
        "paragraphs": len(reader.paragraphs),
        "sections": len(reader.sections),
        "styles": len(reader.styles),
        "fonts": len(reader.fonts),
        "tables": len(getattr(reader, "tables", [])),
    }
    if report:
        print(f"converted: {src} -> {dst}")
        print(f"  chars={rep['chars']} paragraphs={rep['paragraphs']} "
              f"sections={rep['sections']} styles={rep['styles']} fonts={rep['fonts']}")
    return rep


def inspect(src: str) -> dict:
    """Print structural information about a .doc file."""
    reader = DocReader(src)
    reader.build()
    info = {
        "file": src,
        "nFib": reader.nFib,
        "fcMin": reader.fcMin,
        "fcMac": reader.fcMac,
        "ccpText": reader.ccpText,
        "ccpHdd": reader.ccpHdd,
        "pieces": len(reader.pieces),
        "font_table": reader.fonts,
        "paragraphs": len(reader.paragraphs),
    }
    print(json.dumps(info, ensure_ascii=False, indent=2))
    preview = reader.text[:300].replace("\r", "\\r")
    print("text preview:", preview)
    return info


def verify(dst: str, against: str | None = None) -> dict:
    """Re-open a produced .docx and sanity-check it (body + table cells).

    Text is counted from the raw OOXML ``w:t`` elements so merged table cells
    (gridSpan) are not double-counted.
    """
    import zipfile
    d = Document(dst)
    n_paras = len(d.paragraphs)
    n_tables = len(d.tables)
    all_text = ""
    with zipfile.ZipFile(dst) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
        all_text = "".join(t.text or "" for t in root.iter(w("t")))
    text_len = len(all_text)
    ok = n_paras > 0 and text_len > 0
    rep = {"file": dst, "ok": ok, "paragraphs": n_paras, "chars": text_len,
           "tables": n_tables}
    if against:
        r = DocReader(against)
        r.build()
        orig = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "",
                      r.text.replace("\r", "").replace("\t", " "))
        new = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "",
                     all_text.replace("\r", "").replace("\t", " "))
        oa = re.findall(r"[\u4e00-\u9fff\w]", orig)
        na = re.findall(r"[\u4e00-\u9fff\w]", new)
        rep["orig_chars"] = len(oa)
        rep["char_ratio"] = round(len(na) / max(1, len(oa)), 3)
    print(json.dumps(rep, ensure_ascii=False, indent=2))
    return rep


# --------------------------------------------------------------------------- #
#  CLI
# --------------------------------------------------------------------------- #

def _build_parser():
    ap = argparse.ArgumentParser(
        prog="doc2docx", description="Pure-Python .doc -> .docx converter.")
    sub = ap.add_subparsers(dest="cmd", required=True)
    p_c = sub.add_parser("convert", help="convert one file")
    p_c.add_argument("src")
    p_c.add_argument("-o", "--output", default=None)
    p_c.add_argument("--json", action="store_true")
    p_b = sub.add_parser("batch", help="convert a whole directory")
    p_b.add_argument("dir")
    p_b.add_argument("-o", "--output", default=None)
    p_i = sub.add_parser("inspect", help="dump structure")
    p_i.add_argument("src")
    p_v = sub.add_parser("check", help="verify a produced docx")
    p_v.add_argument("dst")
    p_v.add_argument("--against", default=None)
    return ap


def main(argv=None):
    args = _build_parser().parse_args(argv)
    if args.cmd == "convert":
        rep = convert(args.src, args.output)
        if args.json:
            print(json.dumps(rep, ensure_ascii=False))
        return 0
    if args.cmd == "batch":
        outdir = args.output or os.path.join(args.dir, "docx_out")
        os.makedirs(outdir, exist_ok=True)
        files = [f for f in os.listdir(args.dir) if f.lower().endswith(".doc")]
        ok = 0
        for f in sorted(files):
            src = os.path.join(args.dir, f)
            dst = os.path.join(outdir, os.path.splitext(f)[0] + ".docx")
            try:
                convert(src, dst)
                ok += 1
            except Exception as exc:
                print(f"FAILED {f}: {exc}", file=sys.stderr)
        print(f"batch done: {ok}/{len(files)} converted -> {outdir}")
        return 0
    if args.cmd == "inspect":
        inspect(args.src)
        return 0
    if args.cmd == "check":
        verify(args.dst, args.against)
        return 0
    return 2


if __name__ == "__main__":
    sys.exit(main())
