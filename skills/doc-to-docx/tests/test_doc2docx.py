#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""test_doc2docx.py — self-verification for the pure-Python .doc -> .docx
converter. Builds a synthetic Word 97 binary .doc in memory (no external
samples or libraries), converts it, then asserts that text, character
formatting, paragraph alignment, page setup and table structure survive.

Run directly::

    python tests/test_doc2docx.py

or with pytest::

    pytest tests/test_doc2docx.py -v
"""

from __future__ import annotations

import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))
sys.path.insert(0, os.path.dirname(__file__))

import make_doc            # noqa: E402
import doc2docx           # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

FAILED = 0


def check(name: str, cond: bool):
    global FAILED
    if cond:
        print(f"  ok  {name}")
    else:
        FAILED += 1
        print(f"FAIL  {name}")


def run():
    print("doc2docx self-test")
    tmp = tempfile.mkdtemp(prefix="doc2docx_test_")
    src = os.path.join(tmp, "sample.doc")
    dst = os.path.join(tmp, "sample.docx")
    with open(src, "wb") as fh:
        fh.write(make_doc.build_doc())

    reader = doc2docx.DocReader(src)
    reader.build()

    # ---- text --------------------------------------------------------------
    expected_text = ("你好，世界 Hello World\rSecond paragraph\r"
                     "A\x07B\x07\x07C\x07D\x07\x07")
    check("main text extracted byte-exact", reader.text == expected_text)

    # ---- fonts -------------------------------------------------------------
    check("font table parsed", reader.fonts.get(0) == "Times New Roman"
          and reader.fonts.get(1) == "宋体")

    # ---- paragraphs --------------------------------------------------------
    paras = reader.paragraphs
    check("paragraph count == 8", len(paras) == 8)
    check("para1 text",
          reader.text[paras[0].cp_start:paras[0].cp_end] == "你好，世界 Hello World\r")
    check("para2 centered (jc=1)",
          any(p.pfmt.jc == 1 for p in paras if "Second" in reader.text[p.cp_start:p.cp_end]))
    check("table rows found", len(reader.tables) == 1)
    if reader.tables:
        rows = reader.tables[0]
        check("table has 2 rows", len(rows) == 2)
        cells0 = rows[0][1]
        if len(cells0) >= 2:
            t0 = [reader.text[c[0].cp_start:c[0].cp_end].replace("\x07", "")
                  for c in cells0]
            check("row1 cells A/B", t0[0] == "A" and t0[1] == "B")

    # ---- sections ----------------------------------------------------------
    sec = reader.sections[0]
    check("page width A4", sec.page_w == 11906)
    check("page height A4", sec.page_h == 16838)
    check("left margin", sec.margin_l == 1800)

    # ---- style default -----------------------------------------------------
    st0 = reader.styles.get(0)
    check("Normal style size 21", bool(st0) and st0.get("cf").size == 21)

    # ---- convert -----------------------------------------------------------
    rep = doc2docx.convert(src, dst, report=False)
    check("conversion produced file", os.path.exists(dst))
    doc = Document(dst)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    check("body text contains para1",
          "你好，世界 Hello World" in all_text)
    check("body text contains para2",
          "Second paragraph" in all_text)
    check("table built", len(doc.tables) == 1)
    if doc.tables:
        t = doc.tables[0]
        check("docx table 2 rows", len(t.rows) == 2)
        if len(t.rows) and len(t.rows[0].cells) >= 2:
            check("docx row1 cells",
                  t.rows[0].cells[0].text == "A" and t.rows[0].cells[1].text == "B")
    # page setup in the docx
    s = doc.sections[0]
    check("docx A4 width", abs(s.page_width - 11906 * 635) < 1000)
    # character formatting: first run of para1 should be bold, size 32
    p1 = None
    for p in doc.paragraphs:
        if p.text.startswith("你好"):
            p1 = p
            break
    if p1 is not None and p1.runs:
        r0 = p1.runs[0]
        check("para1 bold", bool(r0.bold))
        check("para1 size 32", r0.font.size is not None and r0.font.size.pt == 16.0)
        rf = r0._element.rPr.rFonts if r0._element.rPr is not None and r0._element.rPr.rFonts is not None else None
        ea = rf.get(qn("w:eastAsia")) if rf is not None else None
        check("para1 eastAsia 宋体", ea == "宋体")

    # ---- verify ------------------------------------------------------------
    v = doc2docx.verify(dst)
    check("verify passes", v["ok"])

    print()
    if FAILED:
        print(f"{FAILED} check(s) FAILED")
        return 1
    print("all checks passed")
    return 0


if __name__ == "__main__":
    sys.exit(run())
