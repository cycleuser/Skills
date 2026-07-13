# -*- coding: utf-8 -*-
"""从 docx 提取 + PDF 参考校对 → 内容润色 → 结构化 HTML"""
import re, json, os
from collections import OrderedDict
from lxml import etree
from docx import Document

SRC = '/Users/fred/Documents/GitHub/Current/课程学习/无人机/无人机驾驶员航空知识手册.docx'
IMG_DIR = '/Users/fred/Documents/GitHub/cycleuser/GangDan/drone_images'
OUT_JSON = '/Users/fred/Documents/GitHub/cycleuser/GangDan/knowledge/drone_handbook_final.json'
os.makedirs(os.path.dirname(OUT_JSON), exist_ok=True)

doc = Document(SRC)

# ---- 1. Image mapping ----
img_map = {}
for rel_id, rel in doc.part.rels.items():
    if "image" in rel.reltype:
        fname = os.path.basename(rel.target_ref)
        img_map[rel_id] = f'images/drone/{fname}'

# ---- 2. Extract content sequence ----
body = doc.element.body
content_seq = []

BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
IMAGEDATA = '{urn:schemas-microsoft-com:vml}imagedata'
IMGID = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
T_TAG = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'

def get_images(elem):
    imgs = []
    for blip in elem.findall('.//' + BLIP):
        rid = blip.get(EMBED)
        if rid and rid in img_map:
            imgs.append(img_map[rid])
    for vml in elem.findall('.//' + IMAGEDATA):
        rid = vml.get(IMGID)
        if rid and rid in img_map:
            imgs.append(img_map[rid])
    return imgs

def clean_text(t):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', t).strip()

for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        text = ''.join(run.text or '' for run in child.findall('.//' + T_TAG))
        text = clean_text(text)
        for path in get_images(child):
            content_seq.append(('image', path))
        if text:
            content_seq.append(('text', text))
    elif tag == 'tbl':
        tbl_idx = len([c for c in content_seq if c[0] == 'table'])
        rows = []
        for row in doc.tables[tbl_idx].rows:
            cells = [' '.join(p.text.strip() for p in c.paragraphs if p.text and p.text.strip()) for c in row.cells]
            rows.append('<tr>' + ''.join('<td>' + (c or '') + '</td>' for c in cells) + '</tr>')
        content_seq.append(('table', '<table>' + ''.join(rows) + '</table>'))

# ---- 3. Content refinement ----
# Run through and merge caption text that follows images
caption_pat = re.compile(r'^(图|表)\s*\d+[\.\d]*\s+')
bullet_pat = re.compile(r'^\(?\s*\d+\s*[\)）]')
subtitle_pat = re.compile(r'^[（(]\s*[一二三四五六七八九十\d]+\s*[）)]')
orphan_pat = re.compile(r'^无人机驾驶员航空知识手册$|^第\s*\d+\s*章\s*$')
skip_pat = re.compile(r'^(图书在版|ISBN|中国版本|CIP|责任编辑|版权所有|印张|开本|印数|定价|字数|版次|印次)')

refined = []  # (type, data)
i = 0
while i < len(content_seq):
    typ, data = content_seq[i]
    
    if typ == 'image':
        # Check if next item is a short text that's a caption
        if i+1 < len(content_seq) and content_seq[i+1][0] == 'text':
            next_t = content_seq[i+1][1]
            if (caption_pat.match(next_t) or
                len(next_t) < 40 and ('图' in next_t or '表' in next_t or next_t.startswith('无人机驾驶员'))) and len(next_t) < 60:
                # Merge: image + caption
                fig_html = f'<figure><img src="{data}" style="max-width:100%;height:auto;display:block;margin:10px auto 4px;border-radius:4px" loading="lazy" alt="">'
                fig_html += f'<figcaption style="text-align:center;font-size:13px;color:var(--muted);margin-bottom:12px">{next_t}</figcaption></figure>'
                refined.append(('html', fig_html))
                i += 2
                continue
        # Image without caption
        refined.append(('image', data))
        i += 1
    
    elif typ == 'text':
        t = data
        # Skip orphans
        if orphan_pat.match(t) and len(t) < 30:
            i += 1
            continue
        if skip_pat.match(t):
            i += 1
            continue
        # Sub-bullets -> bullet HTML
        if bullet_pat.match(t) and len(t) > 5:
            # The text WAS merged into a previous section; clean up
            t2 = re.sub(r'^\s*\(?\s*\d+\s*[\)）]\s*', '', t)
            refined.append(('html', f'<li style="margin-left:24px;list-style-type:decimal">{t2}</li>'))
            i += 1
            continue
        if subtitle_pat.match(t) and len(t) < 40:
            refined.append(('html', f'<h4 style="font-weight:700;margin:14px 0 6px;color:var(--text)">{t}</h4>'))
            i += 1
            continue
        # Plain text
        refined.append(('text', data))
        i += 1
    
    elif typ == 'table':
        refined.append(('html', f'<div class="tbl-wrap">{data}</div>'))
        i += 1

# ---- 4. Chapter segmentation ----
chap_pat = re.compile(r'第\s*(\d+)\s*章\s*(.+)')
sec_pat  = re.compile(r'^(\d+[\.\d]*)\s+(.+)')

# Build text-only lookup for chapter detection
chapter_starts = []
for i, (typ, data) in enumerate(refined):
    if typ not in ('text', 'html'): continue
    x = data if typ == 'text' else re.sub(r'<[^>]+>', '', data)
    m = chap_pat.match(x)
    if not m: continue
    num = int(m.group(1))
    for j in range(i+1, min(i+41, len(refined))):
        t2 = refined[j][1]
        if refined[j][0] in ('text', 'html'):
            t2 = t2 if refined[j][0] == 'text' else re.sub(r'<[^>]+>', '', t2)
            if sec_pat.match(t2) and len(t2) > 10:
                if not any(s[1] == num for s in chapter_starts):
                    chapter_starts.append((i, num, m.group(2).strip()))
                break

# ---- 5. Build sections ----
chapters = []
for idx, (start, num, title) in enumerate(chapter_starts):
    end = chapter_starts[idx+1][0] if idx+1 < len(chapter_starts) else len(refined)
    
    sections = OrderedDict()
    cur = "__intro__"
    buf = ''
    imgs = 0
    tbls = 0
    
    for i in range(start, end):
        typ, data = refined[i]
        if typ in ('text', 'html'):
            t = data if typ == 'text' else re.sub(r'<[^>]+>', '', data).strip()
            # New section?
            m = sec_pat.match(t)
            if m and len(t) < 60 and i > start:
                if buf.strip():
                    sections[cur] = [buf.strip(), imgs, tbls]
                cur, buf, imgs, tbls = t, '', 0, 0
                continue
            # Append to buffer
            if typ == 'text':
                buf += '<p>' + data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') + '</p>\n'
            else:
                buf += data + '\n'
                if '<figure>' in data or '<img' in data: imgs += 1
                if '<table' in data: tbls += 1
        elif typ == 'image':
            buf += f'<img src="{data}" style="max-width:100%;height:auto;display:block;margin:10px auto;border-radius:4px" loading="lazy" alt="">\n'
            imgs += 1
    
    if buf.strip():
        sections[cur] = [buf.strip(), imgs, tbls]
    
    # Clean and build
    sec_list = []
    for s_title, (s_html, s_imgs, s_tbls) in sections.items():
        h = re.sub(r'(<p>\s*</p>\n*)+', '', s_html).strip()
        if len(re.sub(r'<[^>]+>', '', h)) < 15 and s_imgs == 0 and s_tbls == 0:
            continue
        sec_list.append({'title': s_title, 'html': h, 'images': s_imgs, 'tables': s_tbls})
    
    if sec_list:
        chapters.append({
            'id': num, 'title': f'第{num}章 {title}',
            'sections': sec_list,
            'total_images': sum(s['images'] for s in sec_list),
            'total_tables': sum(s['tables'] for s in sec_list)
        })

# ---- 6. Polish text: fix OCR artifacts, normalize whitespace ----
def polish_text(html_text):
    """Fix common OCR issues from poor-quality docx"""
    # Remove extra spaces inside words
    t = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', html_text)
    # Remove spaces before Chinese punctuation
    t = re.sub(r'\s+([，。、；：！？）】])', r'\1', t)
    # Normalize English acronyms: "A D" -> "AD"
    # t = re.sub(r'\b([A-Z])\s+([A-Z])\b', r'\1\2', t)
    # Fix doubled symbols
    t = re.sub(r'——+', '——', t)
    # Fix space before newline
    t = re.sub(r'\s+\n', '\n', t)
    # Fix multiple newlines
    t = re.sub(r'\n{4,}', '\n\n\n', t)
    return t

for ch in chapters:
    ch['title'] = polish_text(ch['title'])
    for sec in ch['sections']:
        sec['title'] = polish_text(sec['title'])
        sec['html'] = polish_text(sec['html'])

with open(OUT_JSON, 'w', encoding='utf-8') as f:
    json.dump(chapters, f, ensure_ascii=False, indent=2)

print(f"Chapters: {len(chapters)}")
for ch in chapters:
    print(f"  第{ch['id']}章: {len(ch['sections'])} sections, {ch['total_images']} images, {ch['total_tables']} tables")
print(f"\nJSON: {os.path.getsize(OUT_JSON)/1024:.0f} KB")

# Verify caption wrapping
sample = chapters[0]['sections'][1]['html']
fig_count = sample.count('<figure>')
img_count = sample.count('<img')
print(f"\nSample section 1.1: {fig_count} figures, {img_count} images")
if fig_count > 0:
    import re as _re
    m = _re.search('<figure>.*?</figure>', sample, _re.DOTALL)
    if m: print(m.group()[:300])
