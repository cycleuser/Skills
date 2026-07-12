#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fill_docx_template.py — Template-preserving Word form filler (TianBiao skill).

Core idea: open the ORIGINAL template (converted to .docx), fill text into the
correct cells while keeping the original fonts/borders/merges. Never rebuild tables.

Requires: python-docx  (pip install python-docx)
Convert legacy templates first with LibreOffice:
    soffice --headless --convert-to docx --outdir <dir> <template.doc>

Reusable helpers:
    vcells(table, ri)                  -> visible cells of a row (merge-aware)
    set_cell_value(cell, text)         -> replace a cell's text, keep font, drop sample color/highlight
    set_multi_paragraph_cell(cell, paras, style_idx) -> fill a big analysis cell with heading/body paragraphs
    clone_para(style_p, text)          -> clone a styled paragraph element with new text
    make_detail_table(headers, rows, base_rpr, sz='18') -> compact roster table (tblHeader + cantSplit)
    verify(path)                       -> print table/row/font summary for a filled doc

This file is a library of patterns; adapt the mapping section to each concrete template.
"""
import copy
import docx
from docx.table import _Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- visible-cell access (merge-aware) ----------
def vcells(table, ri):
    """Return VISIBLE cells of row `ri` left-to-right (merged cells collapsed).
    Use this instead of table.cell(r,c) which expands merges into a logical grid."""
    tr = table.rows[ri]._tr
    return [_Cell(tc, table) for tc in tr.findall(qn('w:tc'))]


# ---------- font-preserving cell fill ----------
def _strip_sample_marks(rpr_el):
    if rpr_el is None:
        return
    for tag in ('w:color', 'w:highlight'):
        e = rpr_el.find(qn(tag))
        if e is not None:
            rpr_el.remove(e)


def set_cell_value(cell, text):
    """Replace a cell's text, keeping the first run's formatting.
    Removes sample red color and yellow highlight commonly used in templates."""
    paras = cell.paragraphs
    p0 = paras[0]
    for p in paras[1:]:
        p._p.getparent().remove(p._p)
    runs = p0.runs
    if not runs:
        run = p0.add_run(text)
    else:
        run = runs[0]
        run.text = text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    _strip_sample_marks(run._element.find(qn('w:rPr')))
    pPr = p0._p.find(qn('w:pPr'))
    if pPr is not None:
        _strip_sample_marks(pPr.find(qn('w:rPr')))


# ---------- styled-paragraph cloning ----------
def _clean_pPr(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is not None:
        _strip_sample_marks(pPr.find(qn('w:rPr')))


def _run_rpr(src_p):
    r = src_p.find(qn('w:r'))
    if r is not None:
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            rp = copy.deepcopy(rPr)
            _strip_sample_marks(rp)
            return rp
    return None


def _mk_run(rpr, text):
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def clone_para(style_p_el, text):
    """Deep-copy a styled <w:p> element, drop its runs, set new text. Returns new element."""
    newp = copy.deepcopy(style_p_el)
    _clean_pPr(newp)
    rpr = _run_rpr(newp)
    for r in newp.findall(qn('w:r')):
        newp.remove(r)
    newp.append(_mk_run(rpr, text))
    return newp


def set_multi_paragraph_cell(cell, paragraphs, heading_style_p, body_style_p, indent_chars='200'):
    """Fill a large analysis cell with alternating heading/body paragraphs.
    paragraphs: list of (text, is_heading).
    heading_style_p / body_style_p: sample <w:p> elements to clone formatting from."""
    heading_rpr = _run_rpr(heading_style_p)
    body_rpr = _run_rpr(body_style_p)
    tc = cell._tc
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)
    for text, is_head in paragraphs:
        src = heading_style_p if is_head else body_style_p
        rpr = heading_rpr if is_head else body_rpr
        newp = copy.deepcopy(src)
        _clean_pPr(newp)
        for r in newp.findall(qn('w:r')):
            newp.remove(r)
        # ensure 2-char first-line indent (公文正文)
        pPr = newp.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr'); newp.insert(0, pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind'); pPr.append(ind)
        ind.set(qn('w:firstLineChars'), indent_chars)
        ind.set(qn('w:firstLine'), '480')
        newp.append(_mk_run(rpr, text))
        tc.append(newp)


# ---------- compact detail (roster) tables ----------
def _borders(sz='4'):
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')
        b.append(el)
    return b


def _detail_rpr(base_rpr, sz='18', bold=False):
    if base_rpr is not None:
        rp = copy.deepcopy(base_rpr)
    else:
        rp = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
            rf.set(qn(a), '宋体')
        rp.append(rf)
    for tag in ('w:sz', 'w:szCs', 'w:b'):
        e = rp.find(qn(tag))
        if e is not None:
            rp.remove(e)
    s = OxmlElement('w:sz'); s.set(qn('w:val'), sz); rp.append(s)
    scs = OxmlElement('w:szCs'); scs.set(qn('w:val'), sz); rp.append(scs)
    if bold:
        rp.append(OxmlElement('w:b'))
    return rp


def make_detail_table(headers, rows, base_rpr=None, sz='18'):
    """Build a compact roster table element:
    - small font (sz half-points, default 18 = 9pt)
    - header row repeats on each page (w:tblHeader)
    - every row cannot split across pages (w:cantSplit)
    - tight cell margins
    Returns a <w:tbl> element ready to insert into the body."""
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '0'); tw.set(qn('w:type'), 'auto'); tblPr.append(tw)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); tblPr.append(jc)
    cm = OxmlElement('w:tblCellMar')
    for side, w in (('top', '10'), ('bottom', '10'), ('left', '60'), ('right', '60')):
        el = OxmlElement('w:' + side); el.set(qn('w:w'), w); el.set(qn('w:type'), 'dxa'); cm.append(el)
    tblPr.append(cm)
    tblPr.append(_borders())
    tbl.append(tblPr)
    grid = OxmlElement('w:tblGrid')
    for _ in headers:
        grid.append(OxmlElement('w:gridCol'))
    tbl.append(grid)

    def row(cells, header=False):
        tr = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trPr.append(OxmlElement('w:cantSplit'))
        if header:
            trPr.append(OxmlElement('w:tblHeader'))
        tr.append(trPr)
        for txt in cells:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcw = OxmlElement('w:tcW'); tcw.set(qn('w:w'), '0'); tcw.set(qn('w:type'), 'auto'); tcPr.append(tcw)
            vA = OxmlElement('w:vAlign'); vA.set(qn('w:val'), 'center'); tcPr.append(vA)
            tc.append(tcPr)
            p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
            j2 = OxmlElement('w:jc'); j2.set(qn('w:val'), 'center'); pPr.append(j2)
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
            sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
            pPr.append(sp)
            p.append(pPr)
            r = OxmlElement('w:r')
            r.append(_detail_rpr(base_rpr, sz=sz, bold=header))
            t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = str(txt)
            r.append(t); p.append(r); tc.append(p); tr.append(tc)
        return tr

    tbl.append(row(headers, header=True))
    for rr in rows:
        tbl.append(row(rr))
    return tbl


def cell_font_rpr(cell):
    """Extract a clean deep-copied rPr from a template cell's first run (for detail tables)."""
    for p in cell.paragraphs:
        for r in p.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rp = copy.deepcopy(rPr)
                _strip_sample_marks(rp)
                return rp
    return None


def clone_table_row(table, template_ri):
    """Return a deep copy of an existing <w:tr> to append as a new data row."""
    return copy.deepcopy(table.rows[template_ri]._tr)


# ---------- verification ----------
def verify(path):
    """Print a structural summary to sanity-check a filled document."""
    d = docx.Document(path)
    print(f"[verify] {path}: tables={len(d.tables)} paragraphs={len(d.paragraphs)}")
    for ti, t in enumerate(d.tables):
        tblPr = t._tbl.tblPr
        has_border = tblPr.find(qn('w:tblBorders')) is not None
        head = ' '.join(c.text.strip() for c in t.rows[0].cells)[:40] if t.rows else ''
        print(f"  table{ti}: rows={len(t.rows)} cols={len(t.columns)} borders={has_border} head={head!r}")
    # flag unresolved format placeholders
    full = '\n'.join(p.text for p in d.paragraphs)
    full += '\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)
    for ph in ('%d', '%.1f', '%.2f', '%s', '%%'):
        if ph in full:
            print(f"  WARNING: unresolved placeholder {ph!r} found in document")


if __name__ == '__main__':
    import sys
    if len(sys.argv) == 3 and sys.argv[1] == 'verify':
        verify(sys.argv[2])
    else:
        print(__doc__)
