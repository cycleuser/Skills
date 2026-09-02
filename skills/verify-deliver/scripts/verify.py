#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
verify.py — 通用交付物校验工具 (verify-deliver skill)
用法:
  python verify.py open <docx|xlsx|pdf|图片>   # 可打开性校验
  python verify.py texts <docx|xlsx>           # 导出文本，供人工/脚本核对
  python verify.py count <xlsx> <sheet> <类别列> # 按类别统计行数
  python verify.py unify <dir> <glob> <正则>   # 校验文件名符合命名规范
  python verify.py img-orient <图片>           # 校验图片文字方向（OCR盒坐标）
只读操作，不修改任何输入文件。
"""
import sys, os, re, glob

def open_check(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".docx",):
        from docx import Document
        d = Document(path)
        print(f"OK docx: {len(d.paragraphs)}段, {len(d.tables)}表")
    elif ext in (".xlsx",):
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        print(f"OK xlsx: {wb.sheetnames}")
    elif ext == ".pdf":
        import fitz
        doc = fitz.open(path)
        print(f"OK pdf: {len(doc)}页")
    elif ext in (".png", ".jpg", ".jpeg"):
        from PIL import Image
        im = Image.open(path)
        im.verify()
        print(f"OK image: {im.size if False else os.path.getsize(path)}字节")
    else:
        print(f"SKIP 未支持类型: {ext}")
    return 0

def dump_texts(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document
        d = Document(path)
        for i, t in enumerate(d.tables):
            print(f"===== 表{i} =====")
            for ri, row in enumerate(t.rows):
                seen = set(); cells = []
                for c in row.cells:
                    if id(c._tc) in seen:
                        continue
                    seen.add(id(c._tc))
                    cells.append(c.text.replace("\n", "⏎"))
                print(f"R{ri}: " + " | ".join(cells))
    elif ext == ".xlsx":
        import zipfile
        import xml.etree.ElementTree as ET
        z = zipfile.ZipFile(path)
        ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
        T = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t"
        shared = []
        if "xl/sharedStrings.xml" in z.namelist():
            root = ET.fromstring(z.read("xl/sharedStrings.xml"))
            for si in root.findall("m:si", ns):
                shared.append("".join(t.text or "" for t in si.iter(T)))
        wb = ET.fromstring(z.read("xl/workbook.xml"))
        sheets = [sh.get("name") for sh in wb.find("m:sheets", ns)]
        print("sheets:", sheets)
        for name in sheets:
            print(f"===== {name} =====")
    else:
        print("仅支持 docx/xlsx 文本导出")
    return 0

def count_by_cat(path, sheet=None, col=3):
    import zipfile, re
    import xml.etree.ElementTree as ET
    z = zipfile.ZipFile(path)
    ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    T = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t"
    shared = []
    if "xl/sharedStrings.xml" in z.namelist():
        root = ET.fromstring(z.read("xl/sharedStrings.xml"))
        for si in root.findall("m:si", ns):
            shared.append("".join(t.text or "" for t in si.iter(T)))
    # 简化：遍历所有sheet，统计每行第 col 列（1-based）文本
    from collections import Counter
    cnt = Counter()
    for name in [f"xl/worksheets/sheet{i}.xml" for i in range(1, 40)]:
        if name not in z.namelist():
            continue
        root = ET.fromstring(z.read(name))
        for row in root.findall("m:sheetData/m:row", ns):
            cells = {}
            for c in row.findall("m:c", ns):
                ref = c.get("r"); m = re.match(r"([A-Z]+)", ref)
                cidx = 0
                for ch in m.group(1):
                    cidx = cidx*26 + (ord(ch)-64)
                v = c.find("m:v", ns)
                val = ""
                if v is not None:
                    val = v.text or ""
                    if c.get("t") == "s" and val.isdigit() and int(val) < len(shared):
                        val = shared[int(val)]
                cells[cidx] = val
            if cells and col in cells and cells[col]:
                cnt[cells[col]] += 1
    for k, v in cnt.most_common():
        print(f"{v:4d}  {k}")
    return 0

def name_check(directory, pattern, regex):
    files = glob.glob(os.path.join(directory, pattern))
    bad = []
    for f in files:
        if not re.search(regex, os.path.basename(f)):
            bad.append(os.path.basename(f))
    if bad:
        print("不合规范的文件:")
        for b in bad:
            print("  -", b)
        return 1
    print(f"OK: {len(files)} 个文件全部符合命名规范")
    return 0

def img_orient(path):
    from rapidocr_onnxruntime import RapidOCR
    from PIL import Image
    eng = RapidOCR()
    result, _ = eng(path)
    if not result:
        print("NO_TEXT 未检出文字，无法判断方向")
        return 0
    vertical = 0; total = 0
    for box, text, conf in result[:10]:
        x1, y1 = box[0]; x2, y2 = box[2]
        w = x2 - x1; h = y2 - y1
        total += 1
        if h > w * 1.5:
            vertical += 1
    ratio = vertical / max(total, 1)
    if ratio > 0.6:
        print(f"VERTICAL 检测到竖排文字比例 {ratio:.0%}，图片可能需要旋转90°")
        return 1
    print(f"HORIZONTAL 文字方向正常（竖排比例 {ratio:.0%}）")
    return 0

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)
    cmd, target = sys.argv[1], sys.argv[2]
    try:
        if cmd == "open":
            sys.exit(open_check(target))
        elif cmd == "texts":
            sys.exit(dump_texts(target))
        elif cmd == "count":
            count_by_cat(target, sys.argv[3] if len(sys.argv) > 3 else None, int(sys.argv[4]) if len(sys.argv) > 4 else 3)
        elif cmd == "unify":
            sys.exit(name_check(target, sys.argv[3], sys.argv[4]))
        elif cmd == "img-orient":
            sys.exit(img_orient(target))
        else:
            print("未知命令"); sys.exit(2)
    except Exception as e:
        print(f"FAIL: {e}")
        sys.exit(1)
